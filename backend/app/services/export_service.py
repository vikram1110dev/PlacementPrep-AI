import io
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from app.models.resume import UserResume

class ExportService:
    
    @staticmethod
    def generate_pdf(resume: UserResume) -> bytes:
        """
        Generates a structured PDF using ReportLab based on JSON data.
        """
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        width, height = letter
        
        data = resume.resume_data
        
        # Simple Header
        c.setFont("Helvetica-Bold", 16)
        c.drawString(72, height - 72, data.get("personal_details", {}).get("full_name", "Unknown Name"))
        
        c.setFont("Helvetica", 10)
        c.drawString(72, height - 90, data.get("personal_details", {}).get("email", ""))
        c.drawString(200, height - 90, data.get("personal_details", {}).get("phone", ""))
        
        # Section: Summary
        c.setFont("Helvetica-Bold", 12)
        c.drawString(72, height - 120, "Professional Summary")
        c.setFont("Helvetica", 10)
        c.drawString(72, height - 135, data.get("professional_summary", "")[:100] + "...")
        
        # Save and return bytes
        c.save()
        buffer.seek(0)
        return buffer.getvalue()

    @staticmethod
    def generate_docx(resume: UserResume) -> bytes:
        """
        Generates a DOCX file using python-docx.
        """
        doc = Document()
        data = resume.resume_data
        
        # Header
        doc.add_heading(data.get("personal_details", {}).get("full_name", "Resume"), 0)
        p = doc.add_paragraph()
        p.add_run(data.get("personal_details", {}).get("email", "")).bold = True
        
        # Summary
        doc.add_heading('Professional Summary', level=1)
        doc.add_paragraph(data.get("professional_summary", ""))
        
        # Experience
        doc.add_heading('Experience', level=1)
        for exp in data.get("experience", []):
            doc.add_heading(f"{exp['role']} at {exp['company']}", level=2)
            doc.add_paragraph(f"{exp['start_date']} - {exp.get('end_date', 'Present')}")
            for desc in exp.get("description", []):
                doc.add_paragraph(desc, style='List Bullet')
                
        # Save to bytes buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
